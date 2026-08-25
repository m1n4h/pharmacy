"""
Document Parser Service
Extracts sales data from CSV, Excel, PDF, DOCX, and Image files.
"""
import re
import csv
import io
from datetime import datetime, date
from typing import List, Dict, Optional


class DocumentParser:

    ALLOWED_EXTENSIONS = {'.csv', '.xlsx', '.xls', '.pdf', '.docx', '.png', '.jpg', '.jpeg'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    @staticmethod
    def parse(file_content: bytes, filename: str) -> Dict:
        ext = '.' + filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        if ext not in DocumentParser.ALLOWED_EXTENSIONS:
            return {"success": False, "error": f"File type '{ext}' not supported. Allowed: {', '.join(DocumentParser.ALLOWED_EXTENSIONS)}"}
        if len(file_content) > DocumentParser.MAX_FILE_SIZE:
            return {"success": False, "error": "File too large. Maximum size is 10MB."}

        try:
            if ext == '.csv':
                return DocumentParser.parse_csv(file_content)
            elif ext in ('.xlsx', '.xls'):
                return DocumentParser.parse_excel(file_content)
            elif ext == '.pdf':
                return DocumentParser.parse_pdf(file_content)
            elif ext == '.docx':
                return DocumentParser.parse_docx(file_content)
            elif ext in ('.png', '.jpg', '.jpeg'):
                return DocumentParser.parse_image(file_content)
        except Exception as e:
            return {"success": False, "error": f"Failed to parse file: {str(e)}"}

    @staticmethod
    def parse_csv(file_content: bytes) -> Dict:
        text = file_content.decode('utf-8-sig', errors='replace')
        reader = csv.DictReader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return {"success": False, "error": "CSV file is empty"}

        headers = [h.strip().lower() for h in rows[0].keys()]
        sales = DocumentParser._extract_sales_from_rows(rows, headers)
        return {"success": True, "data": sales, "row_count": len(rows)}

    @staticmethod
    def parse_excel(file_content: bytes) -> Dict:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_content), read_only=True)
        ws = wb.active
        rows_raw = list(ws.iter_rows(values_only=True))
        if not rows_raw:
            return {"success": False, "error": "Excel file is empty"}

        headers = [str(h).strip().lower() if h else '' for h in rows_raw[0]]
        rows = []
        for row in rows_raw[1:]:
            d = {}
            for i, h in enumerate(headers):
                if h:
                    d[h] = row[i] if i < len(row) else ''
            if any(v for v in d.values()):
                rows.append(d)

        sales = DocumentParser._extract_sales_from_rows(rows, headers)
        return {"success": True, "data": sales, "row_count": len(rows)}

    @staticmethod
    def parse_pdf(file_content: bytes) -> Dict:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""

        if not text.strip():
            return {"success": False, "error": "Could not extract text from PDF"}

        sales = DocumentParser._extract_sales_from_text(text)
        return {"success": True, "data": sales, "raw_text": text[:2000]}

    @staticmethod
    def parse_docx(file_content: bytes) -> Dict:
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        text = "\n".join([p.text for p in doc.paragraphs])

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                text += "\n" + " | ".join(cells)

        if not text.strip():
            return {"success": False, "error": "Could not extract text from DOCX"}

        sales = DocumentParser._extract_sales_from_text(text)
        return {"success": True, "data": sales, "raw_text": text[:2000]}

    @staticmethod
    def parse_image(file_content: bytes) -> Dict:
        from PIL import Image
        img = Image.open(io.BytesIO(file_content))
        text = ""
        try:
            import pytesseract
            text = pytesseract.image_to_string(img)
        except Exception:
            text = str(img.info) if img.info else ""

        if not text.strip():
            return {"success": False, "error": "Could not extract text from image. OCR may not be installed."}

        sales = DocumentParser._extract_sales_from_text(text)
        return {"success": True, "data": sales, "raw_text": text[:2000]}

    @staticmethod
    def _extract_sales_from_rows(rows: List[Dict], headers: List[str]) -> List[Dict]:
        field_map = DocumentParser._map_fields(headers)
        sales_dict = {}

        for row in rows:
            customer = row.get(field_map.get('customer', ''), '') or "Walk-in Customer"
            sale_date = row.get(field_map.get('date', ''), '') or date.today().isoformat()
            med_name = row.get(field_map.get('medicine', ''), '') or row.get(field_map.get('item', ''), '')
            qty_raw = row.get(field_map.get('quantity', ''), '') or row.get(field_map.get('qty', ''), '') or 1
            price_raw = row.get(field_map.get('price', ''), '') or row.get(field_map.get('unit_price', ''), '') or 0

            try:
                qty = int(float(str(qty_raw).replace(',', '')))
            except (ValueError, TypeError):
                qty = 1
            try:
                price = float(str(price_raw).replace(',', '').replace('TSh', '').replace('TZS', '').strip())
            except (ValueError, TypeError):
                price = 0

            if not med_name:
                continue

            key = f"{customer}_{sale_date}"
            if key not in sales_dict:
                sales_dict[key] = {
                    "customer_name": str(customer).strip(),
                    "sale_date": DocumentParser._parse_date(str(sale_date)),
                    "items": [],
                    "total": 0
                }
            sales_dict[key]["items"].append({
                "medicine_name": str(med_name).strip(),
                "quantity": max(qty, 1),
                "price": max(price, 0)
            })
            sales_dict[key]["total"] += qty * price

        return list(sales_dict.values())

    @staticmethod
    def _extract_sales_from_text(text: str) -> List[Dict]:
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        sales = []
        current = None

        customer_patterns = [
            r'(?:customer|patient|name|bill\s*to)\s*[:=]\s*(.+)',
            r'(?:customer|patient)\s+(.+)',
        ]
        date_patterns = [
            r'(?:date|sale\s*date|invoice\s*date)\s*[:=]\s*(\d{1,4}[-/]\d{1,2}[-/]\d{1,4})',
            r'(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
        ]
        item_pattern = re.compile(
            r'(.+?)\s+(?:x\s*)?(\d+)\s*[@x]\s*(?:TSh|TZS)?\s*([\d,\.]+)',
            re.IGNORECASE
        )
        total_pattern = re.compile(
            r'(?:total|grand\s*total|amount|balance)\s*[:=]?\s*(?:TSh|TZS)?\s*([\d,\.]+)',
            re.IGNORECASE
        )

        for line in lines:
            lower = line.lower()

            for pat in customer_patterns:
                m = re.search(pat, lower)
                if m:
                    if current and current["items"]:
                        sales.append(current)
                    current = {
                        "customer_name": m.group(1).strip().title(),
                        "sale_date": date.today().isoformat(),
                        "items": [],
                        "total": 0
                    }
                    break

            for pat in date_patterns:
                m = re.search(pat, lower)
                if m and current:
                    current["sale_date"] = DocumentParser._parse_date(m.group(1))
                    break

            m = item_pattern.search(line)
            if m and current:
                try:
                    qty = int(m.group(2))
                    price = float(m.group(3).replace(',', ''))
                    current["items"].append({
                        "medicine_name": m.group(1).strip().title(),
                        "quantity": qty,
                        "price": price
                    })
                    current["total"] += qty * price
                except ValueError:
                    pass

            m = total_pattern.search(line)
            if m and current and not current["total"]:
                try:
                    current["total"] = float(m.group(1).replace(',', ''))
                except ValueError:
                    pass

        if current and current["items"]:
            sales.append(current)

        if not sales:
            simple_items = []
            for line in lines:
                m = item_pattern.search(line)
                if m:
                    try:
                        simple_items.append({
                            "medicine_name": m.group(1).strip().title(),
                            "quantity": int(m.group(2)),
                            "price": float(m.group(3).replace(',', ''))
                        })
                    except ValueError:
                        pass
            if simple_items:
                total = sum(i["quantity"] * i["price"] for i in simple_items)
                sales.append({
                    "customer_name": "Walk-in Customer",
                    "sale_date": date.today().isoformat(),
                    "items": simple_items,
                    "total": total
                })

        return sales

    @staticmethod
    def _map_fields(headers: List[str]) -> Dict[str, str]:
        mapping = {}
        customer_keys = ['customer', 'patient', 'name', 'bill_to', 'bill to', 'client']
        date_keys = ['date', 'sale_date', 'sale date', 'invoice_date', 'invoice date', 'transaction_date']
        medicine_keys = ['medicine', 'medicine_name', 'medicine name', 'drug', 'product', 'item', 'description', 'product_name']
        quantity_keys = ['quantity', 'qty', 'units', 'count', 'num']
        price_keys = ['price', 'unit_price', 'unit price', 'amount', 'rate', 'cost', 'selling_price']

        for h in headers:
            hl = h.lower().strip()
            if not any(mapping.get(k) for k in ['customer']) and any(k in hl for k in customer_keys):
                mapping['customer'] = h
            elif not any(mapping.get(k) for k in ['date']) and any(k in hl for k in date_keys):
                mapping['date'] = h
            elif not mapping.get('medicine') and any(k in hl for k in medicine_keys):
                mapping['medicine'] = h
            elif not mapping.get('quantity') and any(k in hl for k in quantity_keys):
                mapping['quantity'] = h
            elif not mapping.get('price') and any(k in hl for k in price_keys):
                mapping['price'] = h

        return mapping

    @staticmethod
    def _parse_date(date_str: str) -> str:
        date_str = date_str.strip()
        formats = [
            '%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y', '%d-%m-%Y',
            '%Y/%m/%d', '%d.%m.%Y', '%d %b %Y', '%d %B %Y',
            '%Y-%m-%dT%H:%M:%S',
        ]
        for fmt in formats:
            try:
                return datetime.strptime(date_str, fmt).strftime('%Y-%m-%d')
            except ValueError:
                continue
        return date.today().isoformat()
